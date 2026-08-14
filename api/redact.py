import sys
import os
import io
from flask import Flask, request, jsonify, send_file
import docx

# Add parent directory to sys.path to import the CLI redact_pii engine
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from redact_pii import PIIRedactor

app = Flask(__name__)

def redact_docx_stream(redactor, file_stream, output_stream):
    doc = docx.Document(file_stream)
    
    # 1. Redact body paragraphs
    for p in doc.paragraphs:
        redactor.redact_paragraph_runs(p)
        
    # 2. Redact tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    redactor.redact_paragraph_runs(p)
                    
    # 3. Redact headers and footers (all three types)
    for section in doc.sections:
        headers = [section.header, section.first_page_header, section.even_page_header]
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        
        for h in headers:
            if h:
                for p in h.paragraphs:
                    redactor.redact_paragraph_runs(p)
                for table in h.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                redactor.redact_paragraph_runs(p)
                                
        for f in footers:
            if f:
                for p in f.paragraphs:
                    redactor.redact_paragraph_runs(p)
                for table in f.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                redactor.redact_paragraph_runs(p)
                                
    doc.save(output_stream)

@app.route('/api/redact', methods=['POST'])
def redact_api():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    uploaded_file = request.files['file']
    filename = uploaded_file.filename
    if not filename:
        return jsonify({"error": "Empty filename"}), 400
        
    active_types = request.form.getlist('types')
    if not active_types:
        active_types = ["name", "email", "phone", "company", "address", "ssn", "cc", "dob", "ip"]
        
    redactor = PIIRedactor()
    
    file_bytes = uploaded_file.read()
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == '.docx':
            input_stream = io.BytesIO(file_bytes)
            output_stream = io.BytesIO()
            redact_docx_stream(redactor, input_stream, output_stream)
            output_stream.seek(0)
            
            stats_str = ";".join(f"{k}:{v}" for k, v in redactor.counts.items())
            
            resp = send_file(
                output_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"redacted_{filename}"
            )
            resp.headers['Access-Control-Expose-Headers'] = 'X-Redaction-Stats'
            resp.headers['X-Redaction-Stats'] = stats_str
            return resp
            
        else:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            redacted_text, _ = redactor.redact_text(text_content)
            
            output_stream = io.BytesIO(redacted_text.encode('utf-8'))
            stats_str = ";".join(f"{k}:{v}" for k, v in redactor.counts.items())
            
            resp = send_file(
                output_stream,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f"redacted_{filename}"
            )
            resp.headers['Access-Control-Expose-Headers'] = 'X-Redaction-Stats'
            resp.headers['X-Redaction-Stats'] = stats_str
            return resp
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(port=5000, debug=True)
