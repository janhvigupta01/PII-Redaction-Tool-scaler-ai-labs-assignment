#!/usr/bin/env python3
import os
import re
import random
import docx

class PIIRedactor:
    def __init__(self):
        # Maps to ensure consistent substitution throughout the run
        self.mappings = {
            "name": {},
            "email": {},
            "phone": {},
            "company": {},
            "address": {},
            "ssn": {},
            "cc": {},
            "dob": {},
            "ip": {}
        }
        
        # Fake data pools for replacements
        self.fake_names = [
            "John Doe", "Peter Parker", "Jane Smith", "Alice Johnson", "Bob Miller",
            "Charlie Brown", "Emily Davis", "Frank Wilson", "Grace Taylor", "Henry Anderson",
            "Jack Thomas", "Karen Jackson", "Leo White", "Mia Harris", "Nathan Martin",
            "Olivia Clark", "Paul Lewis", "Quinn Robinson", "Rachel Walker", "Samuel Young",
            "Tina Allen", "Victor King", "Wendy Wright", "Xavier Scott", "Yolanda Adams",
            "Zachary Baker", "Bruce Wayne", "Clark Kent", "Diana Prince", "Barry Allen",
            "Arthur Curry", "Hal Jordan", "Oliver Queen", "Lois Lane", "Selina Kyle"
        ]
        self.fake_companies = [
            "Apex Industries Private Limited", "Stark Enterprises Ltd.", "Wayne Enterprises LLC",
            "Globex Corporation", "Initech LLC", "Acme Corp.", "Umbrella Corporation",
            "Hooli Inc.", "Veer Industries Limited", "Soma Tech Ltd."
        ]
        self.fake_addresses = [
            "123 Innovation Way, Tech District, Bangalore - 560001, Karnataka, India",
            "456 Silicon Boulevard, Suite 100, Hyderabad - 500081, Telangana, India",
            "789 Cyber Park, Sector 62, Noida - 201301, Uttar Pradesh, India",
            "101 Cloud Chamber Road, Chakan, Pune - 410501, Maharashtra, India",
            "202 Maple Avenue, Floor 4, Mumbai - 400001, Maharashtra, India",
            "55 Main Street, Sector 15, Gurugram - 122001, Haryana, India"
        ]
        
        # Keep track of counts to generate unique indexed values if pools run out
        self.counts = {k: 0 for k in self.mappings.keys()}
        
        # Load common English, corporate, and geographic words to filter out in Name detection
        self.common_words = self._load_common_words()

    def _load_common_words(self):
        # A list of lowercase words that should never be treated as names
        words = {
            # Basic English
            "the", "of", "and", "to", "in", "is", "you", "that", "it", "he", "was", "for", "on", "are", "as", "with", 
            "his", "they", "at", "be", "this", "have", "from", "or", "one", "had", "by", "word", "but", "not", "what", 
            "all", "were", "we", "when", "your", "can", "said", "there", "use", "an", "each", "which", "she", "do", 
            "how", "their", "if", "will", "up", "other", "about", "out", "many", "then", "them", "these", "so", "some", 
            "her", "would", "make", "like", "him", "into", "has", "look", "two", "more", "write", "go", "see", "no", 
            "way", "could", "people", "my", "than", "first", "water", "been", "call", "who", "oil", "its", "now", "find", 
            "long", "down", "day", "did", "get", "come", "made", "may", "part", "new", "work", "take", "place", "years", 
            "live", "me", "back", "give", "most", "very", "after", "thing", "our", "just", "name", "good", "sentence", 
            "man", "think", "say", "great", "where", "help", "through", "much", "before", "line", "right", "too", "mean", 
            "any", "same", "tell", "boy", "follow", "came", "want", "show", "also", "around", "farm", "three", "small",
            
            # Months / Days
            "january", "february", "march", "april", "may", "june", "july", "august", "september", "october", "november", "december",
            "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep", "oct", "nov", "dec",
            "monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday",
            
            # Corporate / Legal Terms in Prospectus
            "red", "herring", "prospectus", "company", "board", "director", "executive", "independent", "equity", "shares", 
            "act", "companies", "annexure", "schedule", "table", "page", "section", "chapter", "part", "audit", "financial", 
            "report", "bid", "offer", "investor", "bidders", "bidder", "stock", "exchange", "exchanges", "registrar", 
            "government", "securities", "regulations", "rules", "issue", "capital", "management", "group", "bank", 
            "banker", "bankers", "member", "members", "committee", "meeting", "year", "date", "dates", "history", 
            "business", "industry", "summary", "description", "definition", "definitions", "corporate", "general", 
            "risk", "factors", "legal", "other", "information", "outstanding", "litigation", "material", "development", 
            "developments", "statutory", "contract", "contracts", "declaration", "signed", "signatures", "signature", 
            "promoter", "promoters", "shareholder", "shareholders", "manager", "managers", "compliance", "officer", 
            "auditor", "auditors", "secretary", "secretarial", "advisor", "advisors", "counsel", "counsels", "partner", 
            "partners", "associate", "associates", "office", "offices", "registered", "address", "addresses", 
            "telephone", "phone", "email", "fax", "contact", "website", "websites", "url", "number", "numbers", 
            "time", "times", "place", "places", "state", "states", "country", "countries", "city", "cities", "town", 
            "towns", "village", "villages", "district", "districts", "taluka", "talukas", "road", "street", "avenue", 
            "building", "house", "flat", "plot", "sector", "phase", "zone", "area", "region", "national", "international", 
            "domestic", "global", "public", "private", "listed", "unlisted", "sebi", "roc", "rbi", "ministry", 
            "department", "authority", "authorities", "commission", "commissions", "tribunal", "tribunals", "court", 
            "courts", "high", "supreme", "session", "sessions", "judge", "judges", "justice", "justices", "law", 
            "laws", "circular", "circulars", "notification", "notifications", "order", "orders", "judgment", "judgments", 
            "decree", "decrees", "award", "awards", "resolution", "resolutions", "minutes", "agreement", "agreements", 
            "deed", "deeds", "memorandum", "articles", "association", "abridged", "application", "form", "forms", 
            "subscription", "subscriptions", "allotment", "allotments", "refund", "refunds", "dematerialization", 
            "dematerialised", "trading", "listing", "broker", "brokers", "sponsor", "underwriter", "underwriters", 
            "syndicate", "scsb", "scsbs", "upi", "mechanism", "payment", "payments", "account", "accounts", "escrow", 
            "asba", "applicant", "applicants", "retail", "individual", "institutional", "qualified", "buyer", 
            "buyers", "qib", "qibs", "anchor", "mutual", "fund", "funds", "insurance", "trust", "trusts", "foreign", 
            "portfolio", "fpi", "fpis", "fii", "fiis", "nri", "nris", "oci", "ocis", "huf", "hufs", "body", "bodies", 
            "society", "societies", "resident", "residents", "citizen", "citizens", "alien", "aliens", "annexures",
            "schedules", "tables", "pages", "sections", "chapters", "parts", "audits", "reports", "issues", "meetings",
            "years", "sign", "signed", "signing", "signatories", "signatory", "pune", "mumbai", "maharashtra", "india",
            "delhi", "bangalore", "chennai", "kolkata", "hyderabad", "gujarat", "karnataka", "haryana", "punjab",
            "chakan", "khed", "baner", "taluka", "district", "village", "villages", "road", "street", "avenue",
            "licensed", "license", "licence", "licences", "licensee", "licensees", "licensing", "registration",
            "certificate", "certificates", "certified", "certifying", "certify", "incorporation", "incorporated",
            "liability", "proprietor", "proprietorship", "partnership", "partnerships", "firm", "firms", "proprietors",
            "key", "personnel", "kmp", "promoter", "promoters", "group", "groups", "subsidiary", "subsidiaries",
            "holding", "joint", "venture", "ventures", "collaborator", "collaborators", "collaboration", "collaborations",
            "audited", "unaudited", "consolidated", "standalone", "restated", "statement", "statements", "balance",
            "sheet", "sheets", "profit", "loss", "cash", "flow", "flows", "revenue", "revenues", "income", "incomes",
            # Address and street labels
            "off", "farms", "farm", "near", "opposite", "opp", "behind", "next", "beside", "floor", "wing", "block", 
            "sector", "gate", "gat", "plot", "flat", "building", "house", "tower", "centre", "center", "park", 
            "industrial", "estate", "zone", "colony", "nagar", "chowk", "cross", "lane", "road", "street", "highway", 
            "bypass", "station", "airport", "temple", "church", "school", "college", "hospital", "market", "plaza", 
            "mall", "complex", "society", "apartment", "apartments", "villa", "villas", "residency", "gardens", 
            "garden", "heights", "height", "view", "views", "square", "plaza",
            # Additional corporate / prospectus terms
            "price", "band", "period", "date", "dates", "opening", "closing", "issue", "equity", "share", "shares", 
            "bid", "bids", "offer", "offers", "lot", "lots", "working", "days", "day", "information", "business", 
            "financial", "results", "operation", "operations", "condition", "conditions", "risk", "factor", "factors", 
            "legal", "proceedings", "proceeding", "outstanding", "litigation", "litigations", "material", "development", 
            "developments", "corporate", "structure", "history", "matters", "matter", "management", "discussion", 
            "analysis", "capital", "structure", "industry", "overview", "regulations", "policies", "policy", 
            "government", "authorities", "authority", "licenses", "license", "licence", "licences", "approvals", 
            "approval", "registrations", "registration", "permissions", "permission", "manufacturing", "facilities", 
            "facility", "warehouses", "warehouse", "offices", "office", "branches", "branch", "registered", "corporate", 
            "address", "addresses", "contact", "details", "detail", "telephone", "phone", "email", "fax", "mobile", 
            "cell", "website", "websites", "url", "compliance", "officer", "company", "secretary", "promoter", 
            "promoters", "shareholder", "shareholders", "director", "directors", "executive", "independent", "board", 
            "committee", "committees", "audit", "stakeholders", "relationship", "nomination", "remuneration", 
            "corporate", "social", "responsibility", "csr", "risk", "management"
        }
        return words

    def _is_common_word(self, w):
        w_lower = w.lower()
        if w_lower in self.common_words:
            return True
        # Simple stemming: strip plural 's', verb endings 'ed' or 'ing'
        stem = w_lower
        if stem.endswith('s') and len(stem) > 3:
            stem = stem[:-1]
        elif stem.endswith('ed') and len(stem) > 4:
            stem = stem[:-2]
        elif stem.endswith('ing') and len(stem) > 5:
            stem = stem[:-3]
        if stem in self.common_words:
            return True
        return False

    def is_cc_number(self, num_str):
        # Clean the string to digits only
        cleaned = re.sub(r'\D', '', num_str)
        if not (13 <= len(cleaned) <= 19):
            return False
        
        # Luhn Algorithm check
        total = 0
        reverse_digits = cleaned[::-1]
        for i, digit in enumerate(reverse_digits):
            val = int(digit)
            if i % 2 == 1:
                val *= 2
                if val > 9:
                    val -= 9
            total += val
        return total % 10 == 0

    def get_fake_value(self, pii_type, original_value):
        # Normalize whitespace in the key
        original_value = re.sub(r'\s+', ' ', original_value.strip())
        
        if original_value in self.mappings[pii_type]:
            return self.mappings[pii_type][original_value]
        
        # Generate new fake value
        fake_val = ""
        if pii_type == "name":
            # Assign fake name from pool, or create a synthetic name if pool is exhausted
            idx = self.counts[pii_type]
            if idx < len(self.fake_names):
                fake_val = self.fake_names[idx]
            else:
                fake_val = f"FakePerson_{idx - len(self.fake_names) + 1}"
            self.counts[pii_type] += 1
            
        elif pii_type == "email":
            # Check if we can derive from a fake name
            parts = original_value.split('@')
            user_part = parts[0]
            domain = "example.com"
            
            # See if we already mapped this name or user part
            # E.g. ajay.patil@gmail.com -> ajay.patil -> map to assigned name
            # For simplicity, if we have a name mapped, we can use it, else generic
            fake_val = f"user_{self.counts[pii_type] + 1}@{domain}"
            self.counts[pii_type] += 1
            
        elif pii_type == "phone":
            # Keep country code if present, randomize the rest
            has_plus = original_value.startswith('+')
            cleaned = re.sub(r'\D', '', original_value)
            if len(cleaned) > 10:
                cc = cleaned[:-10]
                body = "".join(str(random.randint(0, 9)) for _ in range(10))
                # Add formatting
                fake_val = f"+{cc} {body[:5]} {body[5:]}" if has_plus else f"{cc} {body[:5]} {body[5:]}"
            else:
                body = "".join(str(random.randint(0, 9)) for _ in range(10))
                fake_val = f"+91 {body[:5]} {body[5:]}"
            self.counts[pii_type] += 1
            
        elif pii_type == "company":
            idx = self.counts[pii_type]
            if idx < len(self.fake_companies):
                fake_val = self.fake_companies[idx]
            else:
                fake_val = f"Apex Enterprises Group_{idx - len(self.fake_companies) + 1} Ltd."
            self.counts[pii_type] += 1
            
        elif pii_type == "address":
            idx = self.counts[pii_type]
            if idx < len(self.fake_addresses):
                fake_val = self.fake_addresses[idx]
            else:
                fake_val = f"{idx + 101} Innovation Park, Phase {idx + 1}, Chakan Industrial Area, Pune - 410501, Maharashtra, India"
            self.counts[pii_type] += 1
            
        elif pii_type == "ssn":
            # Format depends on pattern matched
            if re.match(r'^\d{3}-\d{2}-\d{4}$', original_value):
                # US SSN
                fake_val = f"{random.randint(100, 999)}-{random.randint(10, 99)}-{random.randint(1000, 9999)}"
            elif re.match(r'^[A-Z]{5}\d{4}[A-Z]$', original_value):
                # Indian PAN card
                letters = "".join(random.choices("ABCDEFGHIJKLMNOPQRSTUVWXYZ", k=5))
                digits = "".join(random.choices("0123456789", k=4))
                check = random.choice("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
                fake_val = f"{letters}{digits}{check}"
            elif re.match(r'^\d{4}\s\d{4}\s\d{4}$', original_value):
                # Indian Aadhaar
                fake_val = f"{random.randint(1000, 9999)} {random.randint(1000, 9999)} {random.randint(1000, 9999)}"
            else:
                fake_val = "SSN-REDACTED-9999"
            self.counts[pii_type] += 1
            
        elif pii_type == "cc":
            # Generate random valid Luhn number
            digits = [random.randint(0, 9) for _ in range(15)]
            # Luhn checksum digit calculation
            total = 0
            for i, d in enumerate(reversed(digits)):
                val = d
                if i % 2 == 0:  # Because index is 0-based and reversed, matches odd position from right
                    val *= 2
                    if val > 9:
                        val -= 9
                total += val
            check_digit = (10 - (total % 10)) % 10
            digits.append(check_digit)
            card_num = "".join(map(str, digits))
            # Format as groups of 4
            fake_val = f"{card_num[:4]} {card_num[4:8]} {card_num[8:12]} {card_num[12:]}"
            self.counts[pii_type] += 1
            
        elif pii_type == "dob":
            fake_val = f"{random.randint(1, 28):02d}-{random.randint(1, 12):02d}-{random.randint(1970, 2005)}"
            self.counts[pii_type] += 1
            
        elif pii_type == "ip":
            fake_val = f"192.168.{random.randint(1, 254)}.{random.randint(1, 254)}"
            self.counts[pii_type] += 1

        self.mappings[pii_type][original_value] = fake_val
        return fake_val

    def redact_text(self, text):
        if not text:
            return text, []

        matches_found = []  # List of dicts: {"start": int, "end": int, "type": str, "orig": str, "repl": str}
        
        # Helper to record a match
        def add_match(start, end, pii_type, orig, repl):
            matches_found.append({
                "start": start,
                "end": end,
                "type": pii_type,
                "orig": orig,
                "repl": repl
            })

        # --- 1. Credit Cards (checked with Luhn) ---
        cc_pattern = r'\b(?:\d{4}[-\s]?){3}\d{4}\b|\b\d{13,19}\b'
        for m in re.finditer(cc_pattern, text):
            orig = m.group(0)
            if self.is_cc_number(orig):
                repl = self.get_fake_value("cc", orig)
                add_match(m.start(), m.end(), "cc", orig, repl)

        # --- 2. SSNs (US SSN, Indian Aadhaar, PAN) ---
        # Aadhaar: 12 digits with spaces
        aadhaar_pattern = r'\b\d{4}\s\d{4}\s\d{4}\b'
        # PAN: 5 letters, 4 digits, 1 letter
        pan_pattern = r'\b[A-Z]{5}\d{4}[A-Z]\b'
        # US SSN: 3-2-4 digits
        ssn_pattern = r'\b\d{3}-\d{2}-\d{4}\b'
        
        for p in [ssn_pattern, aadhaar_pattern, pan_pattern]:
            for m in re.finditer(p, text):
                # Verify it doesn't overlap with already matched CCs
                if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                    continue
                orig = m.group(0)
                repl = self.get_fake_value("ssn", orig)
                add_match(m.start(), m.end(), "ssn", orig, repl)

        # --- 3. IP Addresses (IPv4 and basic IPv6) ---
        ipv4_pattern = r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b'
        for m in re.finditer(ipv4_pattern, text):
            orig = m.group(0)
            octets = orig.split('.')
            if all(0 <= int(o) <= 255 for o in octets):
                # Check for overlap
                if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                    continue
                repl = self.get_fake_value("ip", orig)
                add_match(m.start(), m.end(), "ip", orig, repl)

        # --- 4. Email Addresses ---
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for m in re.finditer(email_pattern, text):
            if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                continue
            orig = m.group(0)
            repl = self.get_fake_value("email", orig)
            add_match(m.start(), m.end(), "email", orig, repl)

        # --- 5. Dates of Birth (Dob / Born context) ---
        dob_pattern = r'(?i)\b(?:dob|date\s+of\s+birth|born\s+on|born|birth\s+date)\b\s*(?:is|was|on|as\s+of)?\s*[:=-]?\s*(\b\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}\b|\b\d{4}[-/.]\d{1,2}[-/.]\d{1,2}\b|\b(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},\s+\d{4}\b|\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b)'
        for m in re.finditer(dob_pattern, text):
            # The date value is captured in group 1
            date_str = m.group(1)
            # Find start and end indices of the date_str within the overall match
            date_start = m.start(1)
            date_end = m.end(1)
            
            if any(date_start >= r["start"] and date_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("dob", date_str)
            add_match(date_start, date_end, "dob", date_str, repl)

        # --- 6. Phone Numbers ---
        # Telephone/Phone/Fax/Mobile label-based matches (high recall/precision)
        phone_label_pattern = r'(?i)\b(?:tel|telephone|phone|fax|mobile|cell)\b\s*[:=-]?\s*(\+?\s*(?:91|1)?[-\s]?\(?\d{2,5}\)?[-\s]?\d{3,5}[-\s]?\d{4,6}\b)'
        for m in re.finditer(phone_label_pattern, text):
            phone_str = m.group(1).strip()
            phone_start = m.start(1)
            phone_end = m.end(1)
            
            # Clean up the phone string to check if it's mostly digits
            cleaned = re.sub(r'[\s\(\)+-]', '', phone_str)
            if len(cleaned) < 6: # Too short to be a phone number
                continue
                
            if any(phone_start >= r["start"] and phone_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("phone", phone_str)
            add_match(phone_start, phone_end, "phone", phone_str, repl)

        # Standalone phone patterns (10-digit, etc.)
        standalone_phone = r'\b(?:\+?91[\s\.-]?)?[6-9]\d{9}\b|\+?\s*1[-\s]?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}\b'
        for m in re.finditer(standalone_phone, text):
            if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                continue
            orig = m.group(0)
            repl = self.get_fake_value("phone", orig)
            add_match(m.start(), m.end(), "phone", orig, repl)

        # --- 7. Physical/Mailing Addresses ---
        # Label-based block addresses
        address_label_pattern = r'(?i)\b(?:Registered\s+Office|Corporate\s+Office|Registered\s+and\s+Corporate\s+Office|Office\s+at|Office)\b\s*(?:at|is|is\s+at)?\s*[:=-]?\s*((?:(?!Corporate\s+Office|Registered\s+Office)[^\n;\.]){15,250}?(?:Pune|Mumbai|Bangalore|Hyderabad|India|Maharashtra|Karnataka|Delhi|Haryana)\s*[-–]?\s*\d{3}\s?\d{3}\b(?:(?!Corporate\s+Office|Registered\s+Office)[^\n;\.]){0,50})'
        for m in re.finditer(address_label_pattern, text):
            addr_str = m.group(1).strip()
            addr_start = m.start(1)
            addr_end = m.end(1)
            
            if any(addr_start >= r["start"] and addr_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("address", addr_str)
            add_match(addr_start, addr_end, "address", addr_str, repl)

        # Residing at label address
        residing_pattern = r'(?i)\b(?:residing\s+at|resident\s+of)\b\s*([^\n;\.]{15,150})'
        for m in re.finditer(residing_pattern, text):
            addr_str = m.group(1).strip()
            addr_start = m.start(1)
            addr_end = m.end(1)
            
            if any(addr_start >= r["start"] and addr_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("address", addr_str)
            add_match(addr_start, addr_end, "address", addr_str, repl)

        # Specific standalone address blocks (without labels)
        specific_address_pattern = r'\b(?:11/3,\s*11/4\s*and\s*11/5|201,\s*Tower[-–\s]*2)\b(?:(?!Corporate\s+Office|Registered\s+Office)[^\n;\.]){10,200}?\b(?:Pune|Baner|Chakan)\b(?:(?!Corporate\s+Office|Registered\s+Office)[^\n;\.]){0,50}'
        for m in re.finditer(specific_address_pattern, text):
            addr_str = m.group(0).strip()
            addr_start = m.start(0)
            addr_end = m.end(0)
            
            if any(addr_start >= r["start"] and addr_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("address", addr_str)
            add_match(addr_start, addr_end, "address", addr_str, repl)

        # Gat No/Plot No standalone address blocks
        gat_no_pattern = r'\b(?:Gat\s+No\.|Plot\s+No\.)\s*\d+/\d+(?:,\s*\d+/\d+)*(?:,\s*Village\s+[A-Za-z0-9]+)?\b'
        for m in re.finditer(gat_no_pattern, text):
            addr_str = m.group(0).strip()
            addr_start = m.start(0)
            addr_end = m.end(0)
            
            if any(addr_start >= r["start"] and addr_end <= r["end"] for r in matches_found):
                continue
            repl = self.get_fake_value("address", addr_str)
            add_match(addr_start, addr_end, "address", addr_str, repl)



        # --- 8. Company Names ---
        # Suffix-based company name matching
        company_pattern = r'\b[A-Z0-9][A-Za-z0-9&,\'-]*(?:\s+[A-Z0-9][A-Za-z0-9&,\'-]*)*(?:\s+(?:of|and|for)\s+[A-Z0-9][A-Za-z0-9&,\'-]*)?\s+(?:Private\s+Limited|Pvt\.\s*Ltd\.|Pvt\s+Ltd|Limited|Ltd|Corporation|Corp|LLP|LLC|S\.?A\.?)\b'
        for m in re.finditer(company_pattern, text):
            if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                continue
            orig = m.group(0)
            # Verify the words in the company name aren't all ignore words (e.g. "The Company Limited")
            words = [w.lower() for w in re.findall(r'\b\w+\b', orig)]
            # If the only capitalized word is "Company" or similar, skip
            if len(words) <= 2 and words[0] in ["the", "our", "this", "company"]:
                continue
                
            repl = self.get_fake_value("company", orig)
            add_match(m.start(), m.end(), "company", orig, repl)

        # --- 9. Full Names ---
        # 9a. Title-based Names (Mr., Ms., Dr., Shri, Smt., etc.)
        title_pattern = r'\b(?:Mr\.|Ms\.|Mrs\.|Dr\.|Shri|Smt\.|Late)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b'
        for m in re.finditer(title_pattern, text):
            name_str = m.group(1)
            name_start = m.start(1)
            name_end = m.end(1)
            
            if any(name_start >= r["start"] and name_end <= r["end"] for r in matches_found):
                continue
                
            # Verify words aren't common nouns
            words = name_str.split()
            if any(self._is_common_word(w) for w in words):
                continue
                
            repl = self.get_fake_value("name", name_str)
            add_match(name_start, name_end, "name", name_str, repl)

        # 9b. Capitalized Word Heuristics for Names
        # Match 2 or 3 capitalized words in a row
        cap_pattern = r'\b[A-Z][a-z]{1,14}\s+[A-Z][a-z]{1,14}(?:\s+[A-Z][a-z]{1,14})?\b'
        for m in re.finditer(cap_pattern, text):
            if any(m.start() >= r["start"] and m.end() <= r["end"] for r in matches_found):
                continue
                
            name_candidate = m.group(0)
            words = name_candidate.split()
            
            # Filters:
            # 1. Any word is in the common/corporate words set?
            if any(self._is_common_word(w) for w in words):
                continue
            # 2. First word is a standard lowercase grammatical word in disguise?
            if words[0].lower() in ["our", "the", "this", "that", "they", "from", "when", "here", "with", "both"]:
                continue
                
            repl = self.get_fake_value("name", name_candidate)
            add_match(m.start(), m.end(), "name", name_candidate, repl)

        # Sort all matched PII by length descending to replace larger entities first
        resolved_matches = []
        if matches_found:
            # Sort matches by start position
            matches_found.sort(key=lambda x: x["start"])
            
            # Build the output string by pasting non-PII and PII replacements
            last_idx = 0
            parts = []
            for m in matches_found:
                if m["start"] < last_idx:
                    # Overlapping match (should not happen due to filters)
                    continue
                parts.append(text[last_idx:m["start"]])
                parts.append(m["repl"])
                resolved_matches.append(m)
                last_idx = m["end"]
            parts.append(text[last_idx:])
            redacted_text = "".join(parts)
        else:
            redacted_text = text

        return redacted_text, resolved_matches

    def redact_paragraph_runs(self, paragraph):
        if not paragraph.runs:
            return
        
        # Get full paragraph text
        text = "".join(run.text for run in paragraph.runs)
        if not text:
            return
            
        redacted_text, matches = self.redact_text(text)
        if not matches:
            return
            
        # Perform run replacement using our index mapping algorithm
        # Sort matches in descending order of start index to avoid offsets shift
        matches.sort(key=lambda x: x["start"], reverse=True)
        
        runs = paragraph.runs
        for m in matches:
            start = m["start"]
            end = m["end"]
            repl = m["repl"]
            
            # Recalculate run offsets (runs might have changed)
            run_offsets = []
            current_offset = 0
            for run in runs:
                run_len = len(run.text)
                run_offsets.append((current_offset, current_offset + run_len))
                current_offset += run_len
                
            # Find overlapping runs
            first_run_idx = -1
            last_run_idx = -1
            for idx, (r_start, r_end) in enumerate(run_offsets):
                if r_start <= start < r_end:
                    first_run_idx = idx
                if r_start < end <= r_end:
                    last_run_idx = idx
                    break
            
            if first_run_idx == -1 or last_run_idx == -1:
                # Boundary mismatch fallback
                continue
                
            if first_run_idx == last_run_idx:
                run = runs[first_run_idx]
                r_start, _ = run_offsets[first_run_idx]
                local_start = start - r_start
                local_end = end - r_start
                run.text = run.text[:local_start] + repl + run.text[local_end:]
            else:
                first_run = runs[first_run_idx]
                last_run = runs[last_run_idx]
                
                # First run edit
                r_start_first, _ = run_offsets[first_run_idx]
                local_start = start - r_start_first
                first_run.text = first_run.text[:local_start] + repl
                
                # Last run edit
                r_start_last, _ = run_offsets[last_run_idx]
                local_end = end - r_start_last
                last_run.text = last_run.text[local_end:]
                
                # Clear middle runs
                for idx in range(first_run_idx + 1, last_run_idx):
                    runs[idx].text = ""

    def redact_docx(self, input_path, output_path):
        doc = docx.Document(input_path)
        
        # 1. Redact body paragraphs
        for p in doc.paragraphs:
            self.redact_paragraph_runs(p)
            
        # 2. Redact tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.redact_paragraph_runs(p)
                        
        # 3. Redact headers and footers
        for section in doc.sections:
            header = section.header
            if header:
                for p in header.paragraphs:
                    self.redact_paragraph_runs(p)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self.redact_paragraph_runs(p)
                                
            footer = section.footer
            if footer:
                for p in footer.paragraphs:
                    self.redact_paragraph_runs(p)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self.redact_paragraph_runs(p)
                                
        doc.save(output_path)
        print(f"Redacted DOCX saved successfully to: {output_path}")

    def redact_txt(self, input_path, output_path):
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            
        redacted_lines = []
        for line in lines:
            redacted_line, _ = self.redact_text(line)
            redacted_lines.append(redacted_line)
            
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(redacted_lines)
        print(f"Redacted TXT saved successfully to: {output_path}")

def main():
    import sys
    if len(sys.argv) < 3:
        print("Usage: python redact_pii.py <input_file> <output_file>")
        print("Example: python redact_pii.py prospectus.docx redacted_prospectus.docx")
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' does not exist.")
        sys.exit(1)
        
    redactor = PIIRedactor()
    
    ext = os.path.splitext(input_file)[1].lower()
    if ext == '.docx':
        redactor.redact_docx(input_file, output_file)
    else:
        redactor.redact_txt(input_file, output_file)

if __name__ == "__main__":
    main()
